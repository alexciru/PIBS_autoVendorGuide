from docx import Document
from docx.shared import Inches
from assets_interface import AtlassianAssetsAPI
from docx import Document
from docx.shared import Inches
from assets_interface import AtlassianAssetsAPI
import os
from dotenv import load_dotenv
import tqdm

def replace_text_in_paragraphs(paragraphs, replacements):
    """
    Replace text placeholders in a list of paragraphs.
    
    :param paragraphs: List of docx.paragraph.Paragraph objects
    :param replacements: Dict of {"PLACEHOLDER": "replacement text"}
    """
    for p in paragraphs:
        for placeholder, replacement in replacements.items():
            if placeholder in p.text:
                for run in p.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, replacement)

def insert_images_in_paragraphs(paragraphs, image_replacements):
    """
    Replace image placeholders in a list of paragraphs with actual images.
    
    :param paragraphs: List of docx.paragraph.Paragraph objects
    :param image_replacements: Dict of {"PLACEHOLDER": "path/to/image"}
    """
    for p in paragraphs:
        for placeholder, img_path in image_replacements.items():
            if placeholder in p.text:
                p.text = p.text.replace(placeholder, "")
                run = p.add_run()
                run.add_picture(img_path, width=Inches(2))  # Adjust width as needed

def replace_in_tables(tables, replacements, image_replacements):
    """
    Replace text and image placeholders inside tables.
    
    :param tables: List of docx.table.Table objects
    :param replacements: Dict of text replacements
    :param image_replacements: Dict of image replacements
    """
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_paragraphs(cell.paragraphs, replacements)
                insert_images_in_paragraphs(cell.paragraphs, image_replacements)



def replace_labels_with_images(doc_path, output_path, replacements, image_replacements):
    """
    Replace text and image placeholders in a Word document, including headers, footers, and tables.
    
    :param doc_path: Path to the template Word file (.docx)
    :param output_path: Path to save the modified document
    :param replacements: Dict of text replacements {"PLACEHOLDER": "text"}
    :param image_replacements: Dict of image replacements {"PLACEHOLDER": "image_path"}
    """
    doc = Document(doc_path)

    # Replace in main body
    replace_text_in_paragraphs(doc.paragraphs, replacements)
    insert_images_in_paragraphs(doc.paragraphs, image_replacements)
    replace_in_tables(doc.tables, replacements, image_replacements)

    # Replace in headers and footers
    for section in doc.sections:
        # Header
        replace_text_in_paragraphs(section.header.paragraphs, replacements)
        insert_images_in_paragraphs(section.header.paragraphs, image_replacements)
        replace_in_tables(section.header.tables, replacements, image_replacements)
        
        # Footer
        replace_text_in_paragraphs(section.footer.paragraphs, replacements)
        insert_images_in_paragraphs(section.footer.paragraphs, image_replacements)
        replace_in_tables(section.footer.tables, replacements, image_replacements)

    # Save the modified document
    doc.save(output_path)


def get_pib_information(pib_number):
    """Funtion to retrieve information about a specific PiB."""
    api = AtlassianAssetsAPI(
        email=os.getenv("ATLANTSIA_EMAIL"),     # Replace with your Atlassian email
        api_token=os.getenv("ATLANTSIA_API_TOKEN"),  # Replace with your Atlassian API token
        domain=os.getenv("ATLANTSIA_DOMAIN"),  # Replace with your Atlassian domain
    )

    api.get_workspace_id()
    # GET ALL THE PIBS
    object_ids = api.get_object_ids("objectTypeId = 40")
    # get the pib number from the object id
    object_id = object_ids.where(lambda x: x.endswith(f"-{pib_number}"))


################################################################################

if __name__ == "__main__":
    # Define the variables for the executino
    template_path = "PiB Vendor guide T560 - ALEX_template.docx"
    output_path = "output.docx"
    PiB_Number = "138"
    

    # Initialize the Atlassian Assets API
    # All the secret keys are store in the .env file on the same directory
    # change the åath to .env if need
    load_dotenv(".env")  
    api = AtlassianAssetsAPI(
        email=os.getenv("ATLANTSIA_EMAIL"),    
        api_token=os.getenv("ATLANTSIA_API_TOKEN"),  
        domain=os.getenv("ATLANTSIA_DOMAIN"),  
    )

    # Needed to initialize the API and get the workspace ID
    api.get_workspace_id()

    # Step 1: Retrieve all object IDs of type 40 (PiB)
    # This will return a dictionary with object IDs as keys and object data as values
    # we need to filter the object IDs to find the one that matches the PiB number
    # TODO: Improve this to use a more efficient search to avoid multiple REST API calls
    object_dicc = api.get_object_ids("objectTypeId = 40")
    # get the pib number from the object id
    for id in object_dicc.keys():
        attributes = api.get_object_attributes(id)
        if(attributes["PiB"].endswith(PiB_Number)):
            object_id = id
            break

    pib_data = api.get_object(object_id)
    pib_dic = api.attributes_to_dict(pib_data)
    
    print(f"Object ID: {pib_data['id']}, Name: {pib_data['name']}")
    

    print("-------")
    print("-------")
    print("-------")
    print("-------")

 
    vm_dicc_list = pib_dic["Virtual Machines"]
    print("Virtual Machines:")
    for vm_dicc in vm_dicc_list:
        # TODO add values: name, ip+address, hostname, host memory
        attributes = api.get_object_attributes(id)
        print(vm_dicc['displayValue'] + " - " + vm_dicc['referencedObject']['id'])
        
    
    # Retrieve the Teltonika object that is linked to the PIB
    # We need to obtain the referenced object ID from the PIB attributes and obtain all the values
    # info needed: Model, WiFi SSID, WiFi Password
    print("-------")
    print("-------")
    print("-------")
    print("-------")
    teltonika_dicc = pib_dic["Teltonika"][0]
    teltonika_id = teltonika_dicc['referencedObject']['id']
    teltonika_obj = api.get_object(teltonika_id)
    teltonika_attr = api.attributes_to_dict(teltonika_obj)
    
    print(print(list(teltonika_attr.keys())))
    print(teltonika_attr["Model"][0]["displayValue"])
    print(teltonika_attr["WiFi SSID"][0]["displayValue"])
    print(teltonika_attr["WiFi Password"][0]["displayValue"])


    
    # Retrieve its attributes
    pib_atrb = api.get_object_attributes(object_id)
    teltonika_dicc = api.get_object_attributes(teltonika_id)
    


    # Retrieve the VLANs object
    # This information is not linked to the PIB, so we will not retrieve it here
    # We need to retrieve the all the objects of type 34 (VLANs)
    # and keep the ones that match the Location 
    # info needed: VLAN ID, VLAN Name, VLAN IP Range, Gateway, subMask
    print("-------")
    print("-------")
    vlans_dicc = api.get_object_ids("objectTypeId = 34")
    # pint the number of VLANs found
    print(f"Found {len(vlans_dicc)} VLANs")
    print("VLANs:")
    for vlan_id, vlan_data in vlans_dicc.items():
        vlan_attributes = api.get_object_attributes(vlan_id)
        if "Site" in vlan_attributes:
            if vlan_attributes['Site'] == pib_atrb["Location"]:
                print(f"VLAN ID: {vlan_attributes['ID']}, Name: {vlan_attributes['Description']} - {vlan_attributes['IP Range']} - {vlan_attributes['Gateway']}")




    # Retrieve its attributes
    replacements = {
        "<PIB_NUMBER>": PiB_Number,
        "<CLIENT_NAME>": pib_atrb["Vendor"],
        "<LOC>": pib_atrb["Location"],
        "<AREA>": pib_atrb["Area"],
        "<TELTONIKA>": pib_atrb["Teltonika"],
        "<MODEL>": teltonika_dicc["Model"],
        "<WIFI_SSID>": teltonika_dicc["WiFi SSID"],
        "<WIFI_PASSWORD>": teltonika_dicc["WiFi Password"],
        "<CLIENT_DOMAIN>": "xxxxxxxxxxxxxxx",   # TODO: WHERE TO GET THIS?
        "<CLIENT_USER>": "xxxxxxxxxxxxxxx",     # TODO: WHERE TO GET THIS?
    }

    image_replacements = {
        "{TELTONIKA_OLD}": "/Images/teltonika_old.jpg",
        "{TELTONIKA_NEW}": "/Images/teltonika_old.jpg", # TODO: replace with actual new image path
    }

    replace_labels_with_images(template_path, output_path, replacements, image_replacements)
    print(f"Document saved to {output_path}")